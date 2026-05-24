import sys
import subprocess

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Exam Mentor-GPT Documentation', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Exam Mentor-GPT (also referred to as ElitePrep/ExamMentor) is an AI-powered intelligent preparation system "
        "designed for CSS (Central Superior Services) and PPSC (Provincial Public Service Commissions) competitive exams in Pakistan."
    )
    
    # Core Features
    doc.add_heading('2. Core Features', level=1)
    doc.add_paragraph('Intelligent RAG System: Semantic search over hundreds of indexed study documents.', style='List Bullet')
    doc.add_paragraph('Professional Essay Generator: Generates structured essays on CSS/PPSC topics.', style='List Bullet')
    doc.add_paragraph('Intelligent Answer Evaluator: Rubric-based automated scoring with detailed feedback.', style='List Bullet')
    doc.add_paragraph('Performance Analytics Dashboard: Real-time syllabus completion tracking and weak topic identification.', style='List Bullet')
    doc.add_paragraph('Expert Mentor Chatbot: Interactive Q&A with context-aware AI models.', style='List Bullet')
    
    # Architecture
    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph(
        "The system employs a local-first three-tier architecture:\n"
        "- Frontend: Web-based interface with Glassmorphism UI.\n"
        "- Backend: Flask API serving routing and business logic.\n"
        "- AI & Database: Ollama (Mistral 7B) for LLM, ChromaDB for vector storage, and SQLite for persistence."
    )
    
    # Setup
    doc.add_heading('4. Setup and Usage', level=1)
    doc.add_paragraph(
        "Requirements: Python 3.9+, Ollama, Windows/Linux/macOS.\n"
        "To run the system, set up a virtual environment, install dependencies from requirements.txt, "
        "pull the Mistral model via Ollama, and run the backend using 'start.ps1' or 'python backend/main.py'."
    )
    
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "This system democratizes competitive exam preparation by reducing costs, improving evaluation quality, "
        "and providing robust offline AI support."
    )
    
    file_path = 'e:\\Exam Mentor-GPT\\Exam_Mentor_GPT_Documentation.docx'
    doc.save(file_path)
    print(f"Documentation saved successfully to {file_path}")

if __name__ == '__main__':
    create_doc()
